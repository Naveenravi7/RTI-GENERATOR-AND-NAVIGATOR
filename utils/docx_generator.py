from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def generate_rti_docx(data):
    """
    Generates a Word Document (.docx) for the RTI application.
    'data' dictionary should contain:
    - recipient_pio: The PIO title (e.g., "The Public Information Officer")
    - department_name: The department name
    - department_address: The address of the department
    - applicant_name: Name of the applicant
    - applicant_address: Address for correspondence
    - email: Email address (optional)
    - phone: Phone number (optional)
    - citizenship: "Indian Citizen" or custom
    - subject: Subject description
    - period_relates: Period to which info relates
    - queries: List of strings (point-wise questions)
    - is_bpl: Boolean
    - bpl_proof: String details of BPL proof (if is_bpl is True)
    - payment_mode: e.g., "Indian Postal Order", "Online Payment", etc.
    - payment_details: Transaction ID / Serial Number and Date
    - date: Date of application
    - place: Place of application
    """
    doc = Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title / Format Header
    title = doc.add_paragraph()
    title_run = title.add_run("RTI APPLICATION FORM")
    title_run.bold = True
    title_run.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Under Section 6(1) of the Right to Information Act, 2005")
    sub_run.italic = True
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Spacer
    
    # To address
    p_to = doc.add_paragraph()
    p_to.add_run("To,\n").bold = True
    p_to.add_run(f"{data.get('recipient_pio', 'The Public Information Officer')}\n")
    p_to.add_run(f"{data.get('department_name', '[Name of Public Authority]')}\n")
    p_to.add_run(f"{data.get('department_address', '[Address of Public Authority]')}\n")
    
    # Subject
    p_sub = doc.add_paragraph()
    p_sub.add_run("Subject: ").bold = True
    p_sub.add_run("Application seeking information under the Right to Information Act, 2005.")
    
    # Table or bullet details
    doc.add_paragraph() # Spacer
    
    # 1. Applicant Name
    p = doc.add_paragraph()
    p.add_run("1. Name of the Applicant: ").bold = True
    p.add_run(data.get('applicant_name', ''))
    
    # 2. Address
    p = doc.add_paragraph()
    p.add_run("2. Address for Correspondence: ").bold = True
    p.add_run(data.get('applicant_address', ''))
    
    # Contact (optional but good)
    if data.get('phone') or data.get('email'):
        p = doc.add_paragraph()
        contact_str = ""
        if data.get('phone'):
            contact_str += f"Phone: {data.get('phone')}"
        if data.get('email'):
            if contact_str:
                contact_str += ", "
            contact_str += f"Email: {data.get('email')}"
        p.add_run("   Contact Details: ").bold = True
        p.add_run(contact_str)
        
    # 3. Citizenship
    p = doc.add_paragraph()
    p.add_run("3. Citizenship Status: ").bold = True
    p.add_run("Indian Citizen (RTI is filed only by Indian Citizens)")
    
    # 4. Particulars of Information Sought
    doc.add_paragraph().add_run("4. Particulars of Information Sought:").bold = True
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.add_run("(a) Subject matter of information: ").italic = True
    p.add_run(data.get('subject', ''))
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.add_run("(b) Period to which the information relates: ").italic = True
    p.add_run(data.get('period_relates', 'Not Applicable / Current'))
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.add_run("(c) Specific Information required:").italic = True
    
    queries = data.get('queries', [])
    for idx, query in enumerate(queries, 1):
        q_p = doc.add_paragraph()
        q_p.paragraph_format.left_indent = Inches(0.5)
        q_p.add_run(f"({idx}) ").bold = True
        q_p.add_run(query)
        
    # 5. BPL status
    p = doc.add_paragraph()
    p.add_run("5. Whether the applicant belongs to Below Poverty Line (BPL) category? ").bold = True
    is_bpl = data.get('is_bpl', False)
    p.add_run("Yes" if is_bpl else "No")
    
    if is_bpl:
        p_bpl = doc.add_paragraph()
        p_bpl.paragraph_format.left_indent = Inches(0.25)
        p_bpl.add_run("BPL Card / Certificate Proof Details: ").italic = True
        p_bpl.add_run(data.get('bpl_proof', '[BPL Certificate Details] (Copy Attached)'))
        
    # 6. Fees details
    p = doc.add_paragraph()
    p.add_run("6. Application Fee Details:").bold = True
    
    p_fee = doc.add_paragraph()
    p_fee.paragraph_format.left_indent = Inches(0.25)
    if is_bpl:
        p_fee.add_run("Application fee is EXEMPTED as the applicant belongs to the Below Poverty Line (BPL) category.")
    else:
        p_fee.add_run(f"Fee Amount: Rs. 10/-\n")
        p_fee.add_run(f"Mode of Payment: {data.get('payment_mode', 'Indian Postal Order (IPO)')}\n")
        p_fee.add_run(f"Instrument / Transaction No. & Date: {data.get('payment_details', '[Provide details, e.g., IPO No. 12F 345678 dated DD/MM/YYYY]')}")
        
    # Declarations
    doc.add_paragraph()
    p_dec = doc.add_paragraph()
    p_dec.add_run("Declaration:\n").bold = True
    p_dec.add_run("I state that the information sought does not fall within the restrictions contained in Section 8 of the RTI Act, 2005, and to the best of my knowledge, it pertains to your office.\n")
    p_dec.add_run("I am a citizen of India.")
    
    doc.add_paragraph() # Spacer
    
    # Signature blocks
    p_sig = doc.add_paragraph()
    p_sig.add_run(f"Date: {data.get('date', '')}\n")
    p_sig.add_run(f"Place: {data.get('place', '')}")
    
    # Align signature to right
    p_sig_right = doc.add_paragraph()
    p_sig_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig_right.add_run("\n\n___________________________________\n").bold = True
    p_sig_right.add_run("Signature of the Applicant\n").bold = True
    p_sig_right.add_run(f"Name: {data.get('applicant_name', '')}")
    
    # Save to byte stream for Streamlit download
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
